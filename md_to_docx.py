from pathlib import Path
import re
import sys

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parent


def set_run_font(run, name="Microsoft YaHei", size=10.5):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)


def add_inline(paragraph, text):
    parts = re.split(r"(\*\*.*?\*\*|`.*?`)", text)
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run()
        if part.startswith("**") and part.endswith("**"):
            run.text = part[2:-2]
            run.bold = True
            set_run_font(run)
        elif part.startswith("`") and part.endswith("`"):
            run.text = part[1:-1]
            set_run_font(run, "Consolas", 9.5)
        else:
            run.text = part
            set_run_font(run)


def is_table_row(line: str) -> bool:
    s = line.strip()
    return s.startswith("|") and s.endswith("|")


def build_table(doc: Document, rows):
    headers = [c.strip() for c in rows[0].strip("|").split("|")]
    data_rows = []
    for row in rows[2:]:
        data_rows.append([c.strip() for c in row.strip("|").split("|")])

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = header
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                set_run_font(r, size=10)
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "D9EAF7")
        tc_pr.append(shd)

    for row_data in data_rows:
        cells = table.add_row().cells
        for j in range(len(headers)):
            text = row_data[j] if j < len(row_data) else ""
            cells[j].text = text
            cells[j].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[j].paragraphs:
                for r in p.runs:
                    set_run_font(r, size=10)


def find_markdown():
    candidates = sorted(ROOT.glob("*.md"))
    for p in candidates:
        if "需求文档" in p.name:
            return p
    raise FileNotFoundError("Cannot find source markdown in workspace")


def main():
    if len(sys.argv) >= 2:
        md_path = Path(sys.argv[1]).resolve()
        out_path = Path(sys.argv[2]).resolve() if len(sys.argv) >= 3 else md_path.with_suffix(".docx")
    else:
        md_path = find_markdown()
        out_path = ROOT / "TGG Shop 产品需求文档_对齐MD.docx"

    lines = md_path.read_text(encoding="utf-8").splitlines()

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(10.5)

    for section in doc.sections:
        section.top_margin = Pt(54)
        section.bottom_margin = Pt(54)
        section.left_margin = Pt(54)
        section.right_margin = Pt(54)

    in_code = False
    code_lines = []
    i = 0
    while i < len(lines):
        s = lines[i].rstrip("\n")
        t = s.strip()

        if in_code:
            if t.startswith("```"):
                p = doc.add_paragraph()
                r = p.add_run("\n".join(code_lines))
                set_run_font(r, "Consolas", 9)
                code_lines = []
                in_code = False
            else:
                code_lines.append(s)
            i += 1
            continue

        if not t:
            doc.add_paragraph("")
            i += 1
            continue

        if t.startswith("```"):
            in_code = True
            i += 1
            continue

        if t == "---":
            p = doc.add_paragraph()
            p.add_run().add_break(WD_BREAK.LINE)
            i += 1
            continue

        heading = re.match(r"^(#{1,6})\s+(.*)$", t)
        if heading:
            level = min(len(heading.group(1)), 9)
            p = doc.add_paragraph(style=f"Heading {level}")
            add_inline(p, heading.group(2))
            i += 1
            continue

        if is_table_row(t):
            rows = []
            while i < len(lines) and is_table_row(lines[i]):
                rows.append(lines[i].strip())
                i += 1
            build_table(doc, rows)
            continue

        if t.startswith(">"):
            p = doc.add_paragraph()
            add_inline(p, t[1:].strip())
            i += 1
            continue

        if t.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, t[2:])
            i += 1
            continue

        if re.match(r"^\d+\.\s+", t):
            p = doc.add_paragraph(style="List Number")
            add_inline(p, re.sub(r"^\d+\.\s+", "", t))
            i += 1
            continue

        p = doc.add_paragraph()
        add_inline(p, t)
        i += 1

    if out_path.exists():
        out_path.unlink()
    doc.save(str(out_path))
    print(str(out_path))


if __name__ == "__main__":
    main()
